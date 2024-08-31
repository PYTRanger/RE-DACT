from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from google.cloud import dialogflow_v2 as dialogflow
import logging
import warnings
import os
from dotenv import load_dotenv  
import numpy as np
from transformers import pipeline
from PyPDF2 import PdfReader
from docx import Document
import re
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from PIL import Image, ImageDraw
import moviepy.editor as mp
from moviepy.editor import VideoFileClip
import face_recognition

load_dotenv()

logging.basicConfig(level=logging.DEBUG)

google_application_credentials = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")
project_id = os.getenv("PROJECT_ID")

if not google_application_credentials or not project_id:
    raise ValueError("Google application credentials or project ID not set")

os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = google_application_credentials

os.environ['TF_CPP_MIN_LOG_LEVEL'] = '3'
os.environ['TF_ENABLE_ONEDNN_OPTS'] = '0'
warnings.filterwarnings("ignore", category=FutureWarning)
warnings.filterwarnings("ignore", category=DeprecationWarning)
logging.getLogger("transformers").setLevel(logging.ERROR)

app = Flask(__name__)
CORS(app)

client = dialogflow.SessionsClient()
nlp = pipeline("ner", model="dbmdz/bert-large-cased-finetuned-conll03-english", tokenizer="dbmdz/bert-large-cased-finetuned-conll03-english")

def pdfxtract(file_path):
    reader = PdfReader(file_path)
    text = [page.extract_text() for page in reader.pages]
    return "\n".join(text)

def docxtract(file_path):
    doc = Document(file_path)
    text = [p.text for p in doc.paragraphs]
    return "\n".join(text)

def extract_text(file_path):
    ext = os.path.splitext(file_path)[1]
    if ext == '.pdf':
        return pdfxtract(file_path)
    elif ext == '.docx':
        return docxtract(file_path)
    else:
        raise ValueError("Unsupported file format")

def extractInfo(text):
    doc = nlp(text)
    names = []
    current_name = []
    for entity in doc:
        if entity['entity'] in ['B-PER', 'I-PER']:
            if entity['word'].startswith("##"):
                current_name.append(entity['word'][2:])
            else:
                if current_name:
                    names.append("".join(current_name))
                    current_name = []
                current_name.append(entity['word'])
        else:
            if current_name:
                names.append("".join(current_name))
                current_name = []
    if current_name:
        names.append("".join(current_name))
    phones = re.findall(r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text)
    addresses = re.findall(r'\d{1,5} \w+ (Street|St|Avenue|Ave|Boulevard|Blvd|Road|Rd|Lane|Ln|Block|Sector)\b', text)
    return {
        'names': names,
        'phones': phones,
        'addresses': addresses
    }

def redact(text, info):
    for name in info['names']:
        text = text.replace(name, '[REDACTED]')
    for phone in info['phones']:
        text = text.replace(phone, '[REDACTED]')
    for address in info['addresses']:
        text = text.replace(address, '[REDACTED]')
    return text

def redactedDocx(text, output_path):
    doc = Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    doc.save(output_path)
    logging.debug(f"DOCX saved: {output_path}")

def redactedPdf(text, output_path):
    c = canvas.Canvas(output_path, pagesize=letter)
    text_object = c.beginText(40, 750)
    lines = text.split('\n')
    for line in lines:
        text_object.textLine(line)
    c.drawText(text_object)
    c.save()
    logging.debug(f"PDF saved: {output_path}")

def censor_faces(image_path, output_path):
    image = face_recognition.load_image_file(image_path)
    face_locations = face_recognition.face_locations(image)

    pil_image = Image.fromarray(image)
    for face_location in face_locations:
        top, right, bottom, left = face_location
        draw = ImageDraw.Draw(pil_image)
        draw.rectangle([left, top, right, bottom], fill="black")

    pil_image.save(output_path)
    logging.debug(f"Faces censored and image saved: {output_path}")


@app.route('/upload-image', methods=['POST'])
def upload_image():
    if 'file' not in request.files:
        return jsonify({'message': 'No file part'})
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'message': 'No selected file'})

    file_ext = os.path.splitext(file.filename)[1]
    if file_ext.lower() not in ['.png', '.jpg', '.jpeg']:
        return jsonify({'message': 'Unsupported file format'})
    
    file_path = f"temp_image{file_ext}"
    file.save(file_path)
    
    return jsonify({'file_path': file_path})

@app.route('/censor-faces', methods=['POST'])
def censor_faces_route():
    req = request.get_json(silent=True, force=True)
    file_path = req.get('file_path', None)

    if file_path is None or not os.path.isfile(file_path):
        return jsonify({'reply': 'Invalid or missing file'})

    try:
        output_path = "censored_image.png"
        censor_faces(file_path, output_path)
        
        return jsonify({'reply': f"Censored image saved as: {output_path}", 'file_url': f'http://localhost:5000/download/{output_path}'})
    
    except Exception as e:
        logging.error(f"Error processing image: {str(e)}")
        return jsonify({'reply': f"Error: {str(e)}"})

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'message': 'No file part'})
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'message': 'No selected file'})
    
    file_ext = os.path.splitext(file.filename)[1]
    file_path = f"temp_file{file_ext}"
    file.save(file_path)
    
    return jsonify({'file_path': file_path})

@app.route('/send', methods=['POST'])
def send():
    req = request.get_json(silent=True, force=True)
    file_path = req.get('file_path', None)
    entity_to_redact = req.get('entity_type', 'names')

    if file_path is None or not os.path.isfile(file_path):
        return jsonify({'reply': 'Invalid or missing file'})

    try:
        if file_path.endswith('.docx') or file_path.endswith('.pdf'):
            text = extract_text(file_path)
            info = extractInfo(text)
            
            if entity_to_redact == 'names':
                info = {'names': info['names'], 'phones': [], 'addresses': []}
            elif entity_to_redact == 'phones':
                info = {'names': [], 'phones': info['phones'], 'addresses': []}
            elif entity_to_redact == 'addresses':
                info = {'names': [], 'phones': [], 'addresses': info['addresses']}
            
            redacted = redact(text, info)
            
            output_ext = '.docx' if file_path.endswith('.docx') else '.pdf'
            output_path = f"redacted_output{output_ext}"
            
            if output_ext == '.docx':
                redactedDocx(redacted, output_path)
            elif output_ext == '.pdf':
                redactedPdf(redacted, output_path)
        else:
            return jsonify({'reply': 'Unsupported file format'})

        return jsonify({'reply': f"File saved as: {output_path}", 'file_url': f'http://localhost:5000/download/{output_path}'})

    except Exception as e:
        logging.error(f"Error processing file: {str(e)}")
        return jsonify({'reply': f"Error: {str(e)}"})

@app.route('/download/<filename>', methods=['GET'])
def download(filename):
    if not os.path.isfile(filename):
        return jsonify({'error': 'File not found'}), 404
    return send_file(filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
