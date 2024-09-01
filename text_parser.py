
import json
import logging
import warnings
import os
import io
import re
# Import additional packages
from transformers import pipeline
from PyPDF2 import PdfReader
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from google.cloud import dialogflow_v2 as dialogflow

os.environ['TF_CPP_MIN_LOG_LEVEL'] = '3'
os.environ['TF_ENABLE_ONEDNN_OPTS'] = '0'
warnings.filterwarnings("ignore", category=FutureWarning)
warnings.filterwarnings("ignore", category=DeprecationWarning)
logging.getLogger("transformers").setLevel(logging.ERROR)

app = Flask(__name__)
CORS(app)
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "redact-names-icmq-29edfc30b1f7.json"

client = dialogflow.SessionsClient()
project_id = 'redact-names-icmq'

nlp = pipeline("ner", model="dbmdz/bert-large-cased-finetuned-conll03-english", tokenizer="dbmdz/bert-large-cased-finetuned-conll03-english")

def pdfxtract(file_path):
    reader = PdfReader(file_path)
    text = [page.extract_text() for page in reader.pages]
    return "\n".join(text)

def docxtract(file_path):
    doc = Document(file_path)
    text = [p.text for p in doc.paragraphs]
    return "\n".join(text)

def extract_text(file_path):
    ext = os.path.splitext(file_path)[1]
    if ext == '.pdf':
        return pdfxtract(file_path)
    elif ext == '.docx':
        return docxtract(file_path)
    else:
        return 'Unsupported file format'
    
def extractInfo(text):
    doc = nlp(text)
    names = []
    current_name = []
    for entity in doc:
        if entity['entity'] in ['B-PER', 'I-PER']:
            if entity['word'].startswith("##"):
                current_name.append(entity['word'][2:])
            else:
                if current_name:
                    names.append("".join(current_name))
                    current_name = []
                current_name.append(entity['word'])
        else:
            if current_name:
                names.append("".join(current_name))
                current_name = []
    if current_name:
        names.append("".join(current_name))
    phones = re.findall(r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text)
    addresses = re.findall(r'\d{1,5} \w+ (Street|St|Avenue|Ave|Boulevard|Blvd|Road|Rd|Lane|Ln|Block|Sector)\b', text)
    return {
        'names': names,
        'phones': phones,
        'addresses': addresses
    }

def redact(text, info):
    for name in info['names']:
        text = text.replace(name, '[REDACTED]')
    for phone in info['phones']:
        text = text.replace(phone, '[REDACTED]')
    for address in info['addresses']:
        text = text.replace(address, '[REDACTED]')
    return text

def redactedDocx(text, output_path):
    doc = Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    doc.save(output_path)

def redactedPdf(text, output_path):
    c = canvas.Canvas(output_path, pagesize=letter)
    text_object = c.beginText(40, 750)  # Starting position
    lines = text.split('\n')
    for line in lines:
        text_object.textLine(line)
    c.drawText(text_object)
    c.save()

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'message': 'No file part'})
    file = request.files['file']
    if file.filename == '':
        return jsonify({'message': 'No selected file'})
    if file:
        file_path = "temp_file" + os.path.splitext(file.filename)[1]
        file.save(file_path)
        entity_to_redact = request.form.get('entity-type', 'names')

        try:
            text = extract_text(file_path)
            info = extractInfo(text)
            if entity_to_redact == 'names':
                info = {'names': info['names'], 'phones': [], 'addresses': []}
            elif entity_to_redact == 'phones':
                info = {'names': [], 'phones': info['phones'], 'addresses': []}
            elif entity_to_redact == 'addresses':
                info = {'names': [], 'phones': [], 'addresses': info['addresses']}
            else:
                pass
            
            redacted = redact(text, info)
            output_path = "redacted_output.docx"
            if file_path.endswith('.docx'):
                redactedDocx(redacted, output_path)
            elif file_path.endswith('.pdf'):
                output_path = "redacted_output.pdf"
                redactedPdf(redacted, output_path)
            else:
                return jsonify({'message': 'Unsupported file format'})
            
            return send_file(output_path, as_attachment=True)

        except Exception as e:
            return jsonify({'message': f"Error: {str(e)}"})

if __name__ == '__main__':
    app.run(debug=True)
